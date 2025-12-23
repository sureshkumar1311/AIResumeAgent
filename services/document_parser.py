"""
Document parser service for extracting text from PDF and Word documents
"""

import io
from typing import Union
import PyPDF2
from docx import Document
import re
import os
import tempfile
import docx2txt
import olefile



class DocumentParser:
    """Service for parsing various document formats"""
    
    async def parse_document(
        self,
        file_content: bytes,
        filename: str
    ) -> str:
        """
        Parse document and extract complete text content
        
        Args:
            file_content: File content as bytes
            filename: Original filename
        
        Returns:
            Complete extracted text content (no truncation)
        """
        try:
            if filename.lower().endswith('.pdf'):
                return await self._parse_pdf(file_content)
            elif filename.lower().endswith('.docx'):
                return await self._parse_docx(file_content)
            elif filename.lower().endswith('.doc'):
                return await self._parse_doc_legacy(file_content)
            else:
                raise ValueError(f"Unsupported file format: {filename}")
        
        except Exception as e:
            raise Exception(f"Failed to parse document {filename}: {str(e)}")
    
    async def _parse_pdf(self, file_content: bytes) -> str:
        """
        Parse PDF document and extract all text
        
        Args:
            file_content: PDF file content as bytes
        
        Returns:
            Complete extracted text from all pages
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            
            # Extract text from all pages
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text:
                    # Add page separator for better context
                    text_content.append(f"--- Page {page_num} ---\n{text}")
            
            # Return complete text without any truncation
            full_text = "\n\n".join(text_content)
            return full_text
        
        except Exception as e:
            raise Exception(f"Failed to parse PDF: {str(e)}")
    
    async def _parse_docx(self, file_content: bytes) -> str:
        """
        Parse DOCX document and extract all text
        
        Args:
            file_content: DOCX file content as bytes
        
        Returns:
            Complete extracted text including tables
        """
        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            text_content = []
            
            # Extract all paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract all text from tables
            for table_num, table in enumerate(doc.tables, 1):
                table_text = [f"\n--- Table {table_num} ---"]
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_text.append(row_text)
                text_content.append("\n".join(table_text))
            
            # Return complete text without any truncation
            full_text = "\n\n".join(text_content)
            return full_text
        
        except Exception as e:
            raise Exception(f"Failed to parse DOCX document: {str(e)}")
    
    async def _parse_doc_legacy(self, file_content: bytes) -> str:
        """
        Parse legacy .doc file using docx2txt
        """
        try:
            # Write to temporary file
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            
            try:
                # Extract text using docx2txt
                text = docx2txt.process(tmp_path)
                
                if not text or not text.strip():
                    raise Exception("No readable text content found in .doc file")
                
                return text.strip()
            
            finally:
                # Clean up temp file
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
        
        except Exception as e:
            raise Exception(
                f"Failed to parse legacy .doc file: {str(e)}. "
                "Please try converting to .docx or PDF format."
            )

    
    async def _parse_word(self, file_content: bytes) -> str:
        """
        Deprecated: Use _parse_docx or _parse_doc_legacy instead
        Kept for backward compatibility
        """
        # Try to detect format and parse accordingly
        file_obj = io.BytesIO(file_content)
        
        # Check if it's OLE (old .doc)
        if olefile.isOleFile(file_obj):
            return await self._parse_doc_legacy(file_content)
        else:
            # Assume it's DOCX (ZIP-based)
            return await self._parse_docx(file_content)